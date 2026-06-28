"""DOCX template renderer for JobOS.

Loads a user-supplied .docx template, validates required {{ token }} placeholders,
fills them with approved content, and writes the result to a new file.

Key design principles:
- The original template file is NEVER modified.
- Paragraph style (w:pPr) is always preserved: fonts, spacing, heading levels,
  list/bullet indentation come from the template, not from the renderer.
- Run formatting (w:rPr: font, bold, italic, size, colour) is copied from the
  first run of each placeholder paragraph to the replacement runs.
- Multiline replacement values (e.g. bullet lists) expand into multiple
  paragraphs, each inheriting the placeholder paragraph's style, rather than
  using soft line-breaks.  This correctly preserves numbered and bulleted lists.
- Placeholders are detected after joining all runs in a paragraph, so placeholders
  split across Word's internal run boundaries are handled correctly.
"""
from __future__ import annotations

import copy
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn as _ns_qn

from jobos.schemas import ApplicationPack
from jobos.document_generator import jd_to_tags, select_bullets, select_skills

# ── Required placeholder lists ─────────────────────────────────────────────────

CV_PLACEHOLDERS: list[str] = [
    "name",
    "contact_line",
    "profile_summary",
    "skills",
    "experience_role",
    "experience_company",
    "experience_dates",
    "experience_bullets",
    "education_section",
    "adjacent_claims_note",
    "do_not_include_note",
]

CL_PLACEHOLDERS: list[str] = [
    "name",
    "contact_line",
    "date",
    "company",
    "role",
    "greeting",
    "opening_paragraph",
    "body_paragraph_1",
    "body_paragraph_2",
    "motivation_paragraph",
    "closing_paragraph",
    "signoff",
]

# ── Exceptions ─────────────────────────────────────────────────────────────────


class DocxGeneratorError(Exception):
    """Base error for DOCX generation failures."""


class TemplateNotFoundError(DocxGeneratorError):
    """Template file does not exist."""


class MissingPlaceholdersError(DocxGeneratorError):
    """Template is missing one or more required {{ token }} placeholders."""


# ── Low-level XML helpers ──────────────────────────────────────────────────────


def _get_first_rPr(para_p):
    """Return a deep copy of the first run's w:rPr from a paragraph element, or None."""
    for child in para_p:
        if child.tag == _ns_qn("w:r"):
            rPr = child.find(_ns_qn("w:rPr"))
            return copy.deepcopy(rPr) if rPr is not None else None
    return None


def _make_run_elem(text: str, rPr=None):
    """Create a <w:r> element containing the given text and optional run properties."""
    r = OxmlElement("w:r")
    if rPr is not None:
        r.append(copy.deepcopy(rPr))
    t = OxmlElement("w:t")
    t.text = text
    if text != text.strip() or not text:
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    return r


def _clear_content_runs(para_p) -> None:
    """Remove all content-bearing child elements from a paragraph element."""
    for child in list(para_p):
        if child.tag in (
            _ns_qn("w:r"),
            _ns_qn("w:hyperlink"),
            _ns_qn("w:ins"),
            _ns_qn("w:del"),
        ):
            para_p.remove(child)


# ── Paragraph replacement ──────────────────────────────────────────────────────


def _overwrite_paragraph(para, new_text: str) -> None:
    """Replace a paragraph's run content with new_text, preserving formatting.

    - Paragraph style (pPr: heading level, list indentation, spacing) is preserved.
    - Run formatting (rPr: font, bold, italic, size) is copied from the first run.
    - Newlines in new_text become soft line-breaks (w:br) within the paragraph.

    For multiline content that should expand into separate paragraphs (e.g. bullet
    lists), use _expand_paragraph instead — _replace_in_doc does this automatically.
    """
    p = para._p
    rPr = _get_first_rPr(p)
    _clear_content_runs(p)

    lines = new_text.split("\n")
    for i, line in enumerate(lines):
        if i > 0:
            br_run = OxmlElement("w:r")
            if rPr is not None:
                br_run.append(copy.deepcopy(rPr))
            br_run.append(OxmlElement("w:br"))
            p.append(br_run)
        p.append(_make_run_elem(line, rPr))


def _expand_paragraph(para, lines: list[str]) -> None:
    """Replace one placeholder paragraph with one new paragraph per line.

    Each new paragraph is a deep clone of the original (preserving pPr: style,
    list numbering, indent) with run formatting (rPr) from the first original run.
    The original placeholder paragraph is removed from the document after expansion.
    """
    parent = para._p.getparent()
    if parent is None:
        return
    idx = list(parent).index(para._p)
    rPr = _get_first_rPr(para._p)

    for i, line in enumerate(lines):
        new_p = copy.deepcopy(para._p)
        _clear_content_runs(new_p)
        new_p.append(_make_run_elem(line, rPr))
        parent.insert(idx + i, new_p)

    parent.remove(para._p)


def _replace_in_paragraph(para, context: dict[str, str]) -> bool:
    """Detect and replace {{ token }} placeholders in one paragraph (in-place).

    Joins all run texts first so that placeholders split across Word's internal
    run boundaries are handled correctly.

    Returns True if any replacement was made.
    Uses _overwrite_paragraph which preserves paragraph and run formatting.
    """
    full_text = "".join(run.text for run in para.runs)
    if "{{" not in full_text:
        return False

    new_text = full_text
    for token, value in context.items():
        new_text = new_text.replace(f"{{{{ {token} }}}}", value)
        new_text = new_text.replace(f"{{{{{token}}}}}", value)  # no-space variant

    if new_text == full_text:
        return False

    _overwrite_paragraph(para, new_text)
    return True


# ── Document-level replacement ─────────────────────────────────────────────────


def _iter_paragraphs(doc: Document):
    """Yield every paragraph in body, table cells, and section headers/footers."""
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in doc.sections:
        yield from section.header.paragraphs
        yield from section.footer.paragraphs


def _replace_in_doc(doc: Document, context: dict[str, str]) -> None:
    """Apply context replacements to every paragraph in the document.

    Uses a two-phase approach:
    1. Collect every paragraph that contains a placeholder and compute
       its replacement text (joining runs first to handle split-run placeholders).
    2. Apply the replacements.  Multiline values (containing \\n) are expanded
       into multiple paragraphs inheriting the placeholder paragraph's style;
       single-line values are replaced in-place.

    The two-phase design is intentional: Phase 1 must not modify the document
    (so iteration over _iter_paragraphs is safe), and Phase 2 may insert/remove
    paragraph elements.
    """
    # Phase 1: collect replacements
    pending: list[tuple] = []
    for para in _iter_paragraphs(doc):
        full_text = "".join(run.text for run in para.runs)
        if "{{" not in full_text:
            continue
        new_text = full_text
        for token, value in context.items():
            new_text = new_text.replace(f"{{{{ {token} }}}}", value)
            new_text = new_text.replace(f"{{{{{token}}}}}", value)
        if new_text != full_text:
            pending.append((para, new_text))

    # Phase 2: apply replacements
    for para, new_text in pending:
        lines = new_text.split("\n")
        if len(lines) > 1:
            _expand_paragraph(para, lines)
        else:
            _overwrite_paragraph(para, new_text)


# ── Template scanning and validation ──────────────────────────────────────────


def _scan_placeholders(doc: Document, candidates: list[str]) -> dict:
    """Scan a document for placeholders and return location information.

    Returns a dict with:
    - 'found_body'  : list of placeholder names found in body paragraphs
    - 'found_table' : list of placeholder names found in table cells
    - 'all_found'   : de-duplicated union (body ∪ table)
    """
    found_body: list[str] = []
    found_table: list[str] = []

    for para in doc.paragraphs:
        text = "".join(run.text for run in para.runs)
        for ph in candidates:
            if f"{{{{ {ph} }}}}" in text and ph not in found_body:
                found_body.append(ph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = "".join(run.text for run in para.runs)
                    for ph in candidates:
                        if f"{{{{ {ph} }}}}" in text and ph not in found_table:
                            found_table.append(ph)

    all_found = list(dict.fromkeys(found_body + found_table))
    return {"found_body": found_body, "found_table": found_table, "all_found": all_found}


def validate_template(
    path: Path,
    required_placeholders: list[str],
) -> tuple[bool, list[str]]:
    """Check that all required {{ token }} placeholders appear in the template.

    Returns (is_valid, list_of_missing_names).
    Handles placeholders split across Word runs by joining run texts.
    """
    path = Path(path)
    doc = Document(str(path))
    result = _scan_placeholders(doc, required_placeholders)
    missing = [p for p in required_placeholders if p not in result["all_found"]]
    return len(missing) == 0, missing


def diagnose_template(
    path: Path,
    required_placeholders: list[str],
) -> dict:
    """Return rich diagnostic information about a template's placeholders.

    Returns a dict with:
    - 'found_body'  : placeholders found in body paragraphs
    - 'found_table' : placeholders found inside table cells
    - 'all_found'   : union of above
    - 'missing'     : required placeholders not found anywhere
    - 'is_valid'    : bool
    """
    path = Path(path)
    doc = Document(str(path))
    result = _scan_placeholders(doc, required_placeholders)
    missing = [p for p in required_placeholders if p not in result["all_found"]]
    return {
        "found_body": result["found_body"],
        "found_table": result["found_table"],
        "all_found": result["all_found"],
        "missing": missing,
        "is_valid": len(missing) == 0,
    }


# ── Context builders ───────────────────────────────────────────────────────────


def _build_cv_context(
    pack: ApplicationPack,
    profile: dict,
    cv_master: dict,
) -> dict[str, str]:
    meta = cv_master.get("meta", {})
    personal = profile.get("personal", {})

    name = meta.get("name") or personal.get("name", "Sumedh Brahmadevara")
    location = meta.get("location") or personal.get("location", "London, UK")
    email = meta.get("email") or personal.get("email", "[ADD EMAIL]")
    linkedin = meta.get("linkedin") or personal.get("linkedin", "[ADD LINKEDIN]")

    ct = pack.cv_tailor
    jd_tags = jd_to_tags(pack.parsed_job)

    roles = cv_master.get("experience", [])
    all_bullets: list[dict] = [b for role in roles for b in role.get("bullets", [])]
    selected_bullets = select_bullets(all_bullets, jd_tags)
    main_skills, adjacent_skills = select_skills(cv_master.get("skills", []), jd_tags)

    # Primary role (first entry only)
    role_data = roles[0] if roles else {}
    exp_role = role_data.get("title", "")
    exp_company = role_data.get("employer", "")
    start = role_data.get("start", "")
    end = role_data.get("end", "Present")
    exp_dates = f"{start} – {end}"

    first_role_ids = {b["id"] for b in role_data.get("bullets", [])}
    first_role_bullets = [b for b in selected_bullets if b["id"] in first_role_ids]
    exp_bullets = "\n".join(f"• {b['text']}" for b in first_role_bullets)

    # Education
    edu_lines: list[str] = []
    for edu in cv_master.get("education", []):
        diss = edu.get("dissertation", {})
        inst = edu.get("institution", "")
        college = edu.get("college", "")
        heading = inst + (f", {college}" if college else "")
        edu_lines.append(
            f"{edu.get('degree', '')} ({edu.get('grade', '')}) — "
            f"{heading} | {edu.get('years', '')}"
        )
        if diss.get("title"):
            prize = diss.get("prize", "")
            edu_lines.append(
                "Dissertation: " + diss["title"] + (f" — {prize}" if prize else "")
            )
        if diss.get("methods"):
            edu_lines.append("Methods: " + ", ".join(diss["methods"]))
    education_section = "\n".join(edu_lines)

    # Adjacent claims
    adj_items: list[str] = list(ct.adjacent_experience)
    for s in adjacent_skills:
        note = s.get("adjacent_note", "")
        if note:
            adj_items.append(f"{s['text']} — {note}")
    adjacent_note = (
        "\n".join(f"• {item}" for item in adj_items)
        if adj_items
        else "No adjacent experience notes for this role."
    )

    do_not = (
        "\n".join(f"• {c}" for c in ct.unsupported_claims)
        if ct.unsupported_claims
        else "No specific exclusions for this role."
    )

    return {
        "name": name,
        "contact_line": f"{location} │ {email} │ {linkedin}",
        "profile_summary": ct.cv_summary_draft,
        "skills": "\n".join(f"• {s}" for s in main_skills),
        "experience_role": exp_role,
        "experience_company": exp_company,
        "experience_dates": exp_dates,
        "experience_bullets": exp_bullets,
        "education_section": education_section,
        "adjacent_claims_note": adjacent_note,
        "do_not_include_note": do_not,
    }


def _build_cl_context(
    pack: ApplicationPack,
    profile: dict,
    cv_master: dict,
) -> dict[str, str]:
    meta = cv_master.get("meta", {})
    personal = profile.get("personal", {})
    cl_paragraphs = cv_master.get("cover_letter_paragraphs", {})

    name = meta.get("name") or personal.get("name", "Sumedh Brahmadevara")
    location = meta.get("location") or personal.get("location", "London, UK")
    email = meta.get("email") or personal.get("email", "[ADD EMAIL]")

    job = pack.parsed_job
    fit = pack.fit_score
    ct = pack.cv_tailor
    jd_tags = jd_to_tags(job)

    def _select(key: str) -> str:
        options = cl_paragraphs.get(key, [])
        if not options:
            return f"[{key} — add paragraphs to cv_master.yaml]"

        def _score(p: dict) -> int:
            tags = set(p.get("tags", []))
            non_general = tags - {"general"}
            return len(non_general & jd_tags) if non_general else 0

        best = max(options, key=lambda p: (_score(p), p.get("id", "")))
        text = best["text"]
        text = text.replace("{role_title}", job.role_title)
        text = text.replace("{company}", job.company)
        text = text.replace("{application_strategy}", fit.application_strategy)
        return text

    opening = _select("opening") + " " + ct.positioning_angle

    return {
        "name": name,
        "contact_line": f"{location} │ {email}",
        "date": datetime.now().strftime("%d %B %Y"),
        "company": job.company,
        "role": job.role_title,
        "greeting": "Dear Hiring Manager,",
        "opening_paragraph": opening,
        "body_paragraph_1": _select("body_credit"),
        "body_paragraph_2": _select("body_cambridge"),
        "motivation_paragraph": _select("body_motivation"),
        "closing_paragraph": _select("closing"),
        "signoff": f"Yours sincerely,\n\n{name}",
    }


# ── Render functions ───────────────────────────────────────────────────────────


def render_cv_docx(
    template_path: Path,
    pack: ApplicationPack,
    profile: dict,
    cv_master: dict,
    out_path: Path,
) -> Path:
    """Render a tailored CV .docx from a user-supplied template.

    Template formatting (fonts, spacing, bullet styles, margins) is preserved.
    The original template is never modified.

    Raises TemplateNotFoundError if the template does not exist.
    Raises MissingPlaceholdersError if required placeholders are absent.
    """
    template_path = Path(template_path)
    if not template_path.exists():
        raise TemplateNotFoundError(f"CV template not found: {template_path}")

    valid, missing = validate_template(template_path, CV_PLACEHOLDERS)
    if not valid:
        raise MissingPlaceholdersError(
            f"CV template missing required placeholders: {', '.join(missing)}"
        )

    doc = Document(str(template_path))
    context = _build_cv_context(pack, profile, cv_master)
    _replace_in_doc(doc, context)

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def render_cover_letter_docx(
    template_path: Path,
    pack: ApplicationPack,
    profile: dict,
    cv_master: dict,
    out_path: Path,
) -> Path:
    """Render a cover letter .docx from a user-supplied template.

    Template formatting is preserved. The original template is never modified.

    Raises TemplateNotFoundError if the template does not exist.
    Raises MissingPlaceholdersError if required placeholders are absent.
    """
    template_path = Path(template_path)
    if not template_path.exists():
        raise TemplateNotFoundError(f"Cover letter template not found: {template_path}")

    valid, missing = validate_template(template_path, CL_PLACEHOLDERS)
    if not valid:
        raise MissingPlaceholdersError(
            f"Cover letter template missing required placeholders: {', '.join(missing)}"
        )

    doc = Document(str(template_path))
    context = _build_cl_context(pack, profile, cv_master)
    _replace_in_doc(doc, context)

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path
