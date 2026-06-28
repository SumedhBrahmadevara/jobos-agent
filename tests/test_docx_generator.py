"""Tests for jobos/docx_generator.py — DOCX template rendering."""
from __future__ import annotations

import io
from pathlib import Path

import pytest
from docx import Document


# ── Shared fixtures ────────────────────────────────────────────────────────────

_PROFILE = {
    "personal": {
        "name": "Sumedh Brahmadevara",
        "location": "London, UK",
        "email": "test@example.com",
        "linkedin": "linkedin.com/in/test",
    },
}

_ADJACENT = {
    "python_experience": {
        "note": "Python: frame carefully.",
        "safe_phrases": ["Building Python capability for investment workflow tools."],
    }
}


def _make_pack():
    """Build a minimal offline ApplicationPack using the full pipeline."""
    import tempfile
    from apply import build_pack

    jd = (
        "Company: Alpha Fund\n"
        "Role: Research Analyst\n"
        "Fundamental credit research and financial modelling required."
    )
    qs = ""
    with tempfile.TemporaryDirectory() as tmp:
        job_file = Path(tmp) / "job.txt"
        q_file = Path(tmp) / "q.txt"
        job_file.write_text(jd, encoding="utf-8")
        q_file.write_text(qs, encoding="utf-8")
        return build_pack(job_file, q_file)


def _make_cv_docx(tmp_path: Path) -> Path:
    """Create a synthetic CV .docx containing all required CV placeholders."""
    from jobos.docx_generator import CV_PLACEHOLDERS

    doc = Document()
    for placeholder in CV_PLACEHOLDERS:
        doc.add_paragraph(f"{{{{ {placeholder} }}}}")
    path = tmp_path / "cv_template.docx"
    doc.save(str(path))
    return path


def _make_cl_docx(tmp_path: Path) -> Path:
    """Create a synthetic cover letter .docx containing all required CL placeholders."""
    from jobos.docx_generator import CL_PLACEHOLDERS

    doc = Document()
    for placeholder in CL_PLACEHOLDERS:
        doc.add_paragraph(f"{{{{ {placeholder} }}}}")
    path = tmp_path / "cl_template.docx"
    doc.save(str(path))
    return path


# ── Exception class tests ──────────────────────────────────────────────────────

def test_exception_hierarchy():
    from jobos.docx_generator import (
        DocxGeneratorError, TemplateNotFoundError, MissingPlaceholdersError,
    )

    assert issubclass(TemplateNotFoundError, DocxGeneratorError)
    assert issubclass(MissingPlaceholdersError, DocxGeneratorError)


# ── Template not found ─────────────────────────────────────────────────────────

def test_render_cv_docx_raises_template_not_found(tmp_path):
    from jobos.docx_generator import render_cv_docx, TemplateNotFoundError

    missing = tmp_path / "does_not_exist.docx"
    pack = _make_pack()
    cv_master = {}
    with pytest.raises(TemplateNotFoundError):
        render_cv_docx(missing, pack, _PROFILE, cv_master, tmp_path / "out.docx")


def test_render_cl_docx_raises_template_not_found(tmp_path):
    from jobos.docx_generator import render_cover_letter_docx, TemplateNotFoundError

    missing = tmp_path / "does_not_exist.docx"
    pack = _make_pack()
    cv_master = {}
    with pytest.raises(TemplateNotFoundError):
        render_cover_letter_docx(missing, pack, _PROFILE, cv_master, tmp_path / "out.docx")


# ── validate_template ──────────────────────────────────────────────────────────

def test_validate_template_passes_complete_cv(tmp_path):
    from jobos.docx_generator import validate_template, CV_PLACEHOLDERS

    path = _make_cv_docx(tmp_path)
    valid, missing = validate_template(path, CV_PLACEHOLDERS)
    assert valid is True
    assert missing == []


def test_validate_template_passes_complete_cl(tmp_path):
    from jobos.docx_generator import validate_template, CL_PLACEHOLDERS

    path = _make_cl_docx(tmp_path)
    valid, missing = validate_template(path, CL_PLACEHOLDERS)
    assert valid is True
    assert missing == []


def test_validate_template_reports_missing_placeholder(tmp_path):
    from jobos.docx_generator import validate_template

    doc = Document()
    doc.add_paragraph("{{ name }}")
    doc.add_paragraph("{{ contact_line }}")
    path = tmp_path / "partial.docx"
    doc.save(str(path))

    valid, missing = validate_template(path, ["name", "contact_line", "profile_summary"])
    assert valid is False
    assert "profile_summary" in missing


def test_validate_template_reports_all_missing(tmp_path):
    from jobos.docx_generator import validate_template, CV_PLACEHOLDERS

    doc = Document()
    doc.add_paragraph("This template has no placeholders.")
    path = tmp_path / "empty.docx"
    doc.save(str(path))

    valid, missing = validate_template(path, CV_PLACEHOLDERS)
    assert valid is False
    assert len(missing) == len(CV_PLACEHOLDERS)


# ── Missing placeholder error from render functions ────────────────────────────

def test_render_cv_raises_missing_placeholders(tmp_path):
    from jobos.docx_generator import render_cv_docx, MissingPlaceholdersError
    from jobos.document_generator import load_cv_master

    doc = Document()
    doc.add_paragraph("{{ name }} only")
    template = tmp_path / "partial_cv.docx"
    doc.save(str(template))

    cv_master = load_cv_master()
    pack = _make_pack()
    with pytest.raises(MissingPlaceholdersError):
        render_cv_docx(template, pack, _PROFILE, cv_master, tmp_path / "out.docx")


def test_render_cl_raises_missing_placeholders(tmp_path):
    from jobos.docx_generator import render_cover_letter_docx, MissingPlaceholdersError
    from jobos.document_generator import load_cv_master

    doc = Document()
    doc.add_paragraph("{{ name }} only")
    template = tmp_path / "partial_cl.docx"
    doc.save(str(template))

    cv_master = load_cv_master()
    pack = _make_pack()
    with pytest.raises(MissingPlaceholdersError):
        render_cover_letter_docx(template, pack, _PROFILE, cv_master, tmp_path / "out.docx")


# ── CV DOCX rendering ──────────────────────────────────────────────────────────

def test_render_cv_docx_produces_output_file(tmp_path):
    from jobos.docx_generator import render_cv_docx
    from jobos.document_generator import load_cv_master

    template = _make_cv_docx(tmp_path)
    cv_master = load_cv_master()
    pack = _make_pack()
    out = tmp_path / "cv_out.docx"
    result = render_cv_docx(template, pack, _PROFILE, cv_master, out)
    assert result == out
    assert out.exists()


def test_render_cv_docx_output_is_valid_docx(tmp_path):
    from jobos.docx_generator import render_cv_docx
    from jobos.document_generator import load_cv_master

    template = _make_cv_docx(tmp_path)
    cv_master = load_cv_master()
    pack = _make_pack()
    out = tmp_path / "cv_out.docx"
    render_cv_docx(template, pack, _PROFILE, cv_master, out)
    # If python-docx can open it, it's a valid DOCX
    doc = Document(str(out))
    assert len(doc.paragraphs) > 0


def test_render_cv_docx_contains_name(tmp_path):
    from jobos.docx_generator import render_cv_docx
    from jobos.document_generator import load_cv_master

    template = _make_cv_docx(tmp_path)
    cv_master = load_cv_master()
    pack = _make_pack()
    out = tmp_path / "cv_out.docx"
    render_cv_docx(template, pack, _PROFILE, cv_master, out)

    doc = Document(str(out))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Sumedh" in full_text


def test_render_cv_docx_name_from_cv_master(tmp_path):
    """cv_master.meta.name takes priority over profile.personal.name."""
    from jobos.docx_generator import render_cv_docx
    from jobos.document_generator import load_cv_master

    template = _make_cv_docx(tmp_path)
    cv_master = load_cv_master()
    pack = _make_pack()
    out = tmp_path / "cv_out.docx"
    render_cv_docx(template, pack, _PROFILE, cv_master, out)

    doc = Document(str(out))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    # cv_master.yaml defines the name
    assert cv_master["meta"]["name"] in full_text


def test_render_cv_docx_no_unfilled_placeholders(tmp_path):
    from jobos.docx_generator import render_cv_docx, CV_PLACEHOLDERS
    from jobos.document_generator import load_cv_master

    template = _make_cv_docx(tmp_path)
    cv_master = load_cv_master()
    pack = _make_pack()
    out = tmp_path / "cv_out.docx"
    render_cv_docx(template, pack, _PROFILE, cv_master, out)

    doc = Document(str(out))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    for placeholder in CV_PLACEHOLDERS:
        assert f"{{{{ {placeholder} }}}}" not in full_text, (
            f"Placeholder '{{{{ {placeholder} }}}}' was not replaced"
        )


def test_render_cv_docx_profile_summary_not_a_forbidden_claim(tmp_path):
    """The profile_summary and skills context must not contain forbidden claim text.

    Forbidden claims legitimately appear in the do_not_include_note advisory section
    (which tells the reviewer what NOT to claim), so we check the specific fields
    that populate positive CV content, not the full document text.
    """
    from jobos.docx_generator import _build_cv_context
    from jobos.document_generator import load_cv_master

    cv_master = load_cv_master()
    pack = _make_pack()
    context = _build_cv_context(pack, _PROFILE, cv_master)

    # These must never appear as positive claims in the CV content fields
    assert "CFA charterholder" not in context["profile_summary"]
    assert "CFA charterholder" not in context["skills"]
    assert "CFA charterholder" not in context["experience_bullets"]
    assert "Production machine learning engineer" not in context["profile_summary"]


def test_render_cv_docx_does_not_modify_template(tmp_path):
    from jobos.docx_generator import render_cv_docx
    from jobos.document_generator import load_cv_master

    template = _make_cv_docx(tmp_path)
    original_bytes = template.read_bytes()

    cv_master = load_cv_master()
    pack = _make_pack()
    out = tmp_path / "cv_out.docx"
    render_cv_docx(template, pack, _PROFILE, cv_master, out)

    assert template.read_bytes() == original_bytes, "Template file was modified"


def test_render_cv_docx_creates_parent_directory(tmp_path):
    from jobos.docx_generator import render_cv_docx
    from jobos.document_generator import load_cv_master

    template = _make_cv_docx(tmp_path)
    cv_master = load_cv_master()
    pack = _make_pack()
    nested_out = tmp_path / "deep" / "nested" / "cv.docx"
    render_cv_docx(template, pack, _PROFILE, cv_master, nested_out)
    assert nested_out.exists()


# ── Cover letter DOCX rendering ────────────────────────────────────────────────

def test_render_cl_docx_produces_output_file(tmp_path):
    from jobos.docx_generator import render_cover_letter_docx
    from jobos.document_generator import load_cv_master

    template = _make_cl_docx(tmp_path)
    cv_master = load_cv_master()
    pack = _make_pack()
    out = tmp_path / "cl_out.docx"
    result = render_cover_letter_docx(template, pack, _PROFILE, cv_master, out)
    assert result == out
    assert out.exists()


def test_render_cl_docx_contains_company_name(tmp_path):
    from jobos.docx_generator import render_cover_letter_docx
    from jobos.document_generator import load_cv_master

    template = _make_cl_docx(tmp_path)
    cv_master = load_cv_master()
    pack = _make_pack()
    out = tmp_path / "cl_out.docx"
    render_cover_letter_docx(template, pack, _PROFILE, cv_master, out)

    doc = Document(str(out))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert pack.parsed_job.company in full_text


def test_render_cl_docx_no_unfilled_placeholders(tmp_path):
    from jobos.docx_generator import render_cover_letter_docx, CL_PLACEHOLDERS
    from jobos.document_generator import load_cv_master

    template = _make_cl_docx(tmp_path)
    cv_master = load_cv_master()
    pack = _make_pack()
    out = tmp_path / "cl_out.docx"
    render_cover_letter_docx(template, pack, _PROFILE, cv_master, out)

    doc = Document(str(out))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    for placeholder in CL_PLACEHOLDERS:
        assert f"{{{{ {placeholder} }}}}" not in full_text, (
            f"Placeholder '{{{{ {placeholder} }}}}' was not replaced"
        )


def test_render_cl_docx_does_not_modify_template(tmp_path):
    from jobos.docx_generator import render_cover_letter_docx
    from jobos.document_generator import load_cv_master

    template = _make_cl_docx(tmp_path)
    original_bytes = template.read_bytes()

    cv_master = load_cv_master()
    pack = _make_pack()
    out = tmp_path / "cl_out.docx"
    render_cover_letter_docx(template, pack, _PROFILE, cv_master, out)

    assert template.read_bytes() == original_bytes, "Template file was modified"


def test_render_cl_docx_contains_greeting(tmp_path):
    from jobos.docx_generator import render_cover_letter_docx
    from jobos.document_generator import load_cv_master

    template = _make_cl_docx(tmp_path)
    cv_master = load_cv_master()
    pack = _make_pack()
    out = tmp_path / "cl_out.docx"
    render_cover_letter_docx(template, pack, _PROFILE, cv_master, out)

    doc = Document(str(out))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Dear Hiring Manager" in full_text


# ── _overwrite_paragraph internals ────────────────────────────────────────────

def test_overwrite_paragraph_replaces_text():
    from jobos.docx_generator import _overwrite_paragraph

    doc = Document()
    para = doc.add_paragraph("{{ name }}")
    _overwrite_paragraph(para, "Alice Smith")
    assert para.text == "Alice Smith"


def test_overwrite_paragraph_handles_newlines():
    from jobos.docx_generator import _overwrite_paragraph

    doc = Document()
    para = doc.add_paragraph("placeholder")
    _overwrite_paragraph(para, "Line one\nLine two")
    # After overwrite the paragraph text joins lines without the soft break
    # (soft breaks are <w:br/> elements, not \n in para.text)
    combined = para.text
    assert "Line one" in combined
    assert "Line two" in combined


# ── Split-run handling ─────────────────────────────────────────────────────────

def test_replace_handles_split_run():
    """_replace_in_paragraph must handle a placeholder split across two runs."""
    from jobos.docx_generator import _replace_in_paragraph

    doc = Document()
    para = doc.add_paragraph()
    # Simulate Word splitting '{{ name }}' into two runs
    para.add_run("{{ na")
    para.add_run("me }}")

    changed = _replace_in_paragraph(para, {"name": "Bob Jones"})
    assert changed is True
    assert "Bob Jones" in para.text


def test_replace_returns_false_when_no_placeholder():
    from jobos.docx_generator import _replace_in_paragraph

    doc = Document()
    para = doc.add_paragraph("No placeholders here.")
    changed = _replace_in_paragraph(para, {"name": "Alice"})
    assert changed is False


# ── Run formatting preservation ────────────────────────────────────────────────

def test_overwrite_preserves_bold_from_first_run():
    """Bold formatting from the first run must survive _overwrite_paragraph."""
    from jobos.docx_generator import _overwrite_paragraph

    doc = Document()
    para = doc.add_paragraph()
    run = para.add_run("{{ name }}")
    run.bold = True
    _overwrite_paragraph(para, "Alice Smith")
    # The new run should be bold
    assert any(r.bold for r in para.runs if r.text.strip())


def test_overwrite_preserves_font_size_from_first_run():
    """Font size from the first run must survive _overwrite_paragraph."""
    from jobos.docx_generator import _overwrite_paragraph
    from docx.shared import Pt

    doc = Document()
    para = doc.add_paragraph()
    run = para.add_run("{{ name }}")
    run.font.size = Pt(14)
    _overwrite_paragraph(para, "Bob Jones")
    new_runs = [r for r in para.runs if r.text.strip()]
    assert any(r.font.size == Pt(14) for r in new_runs)


# ── _expand_paragraph ──────────────────────────────────────────────────────────

def test_expand_paragraph_removes_original(tmp_path):
    """The original placeholder paragraph must be removed after expansion."""
    from jobos.docx_generator import _expand_paragraph

    doc = Document()
    para = doc.add_paragraph("{{ skills }}")
    original_p = para._p
    _expand_paragraph(para, ["• Skill A", "• Skill B"])
    # original <w:p> must no longer be in the document body
    assert original_p not in list(doc.element.body)


def test_expand_paragraph_creates_correct_count(tmp_path):
    """_expand_paragraph must insert exactly as many paragraphs as lines."""
    from jobos.docx_generator import _expand_paragraph

    doc = Document()
    # Add a sentinel paragraph before and after to check insertion position
    doc.add_paragraph("Before")
    para = doc.add_paragraph("{{ bullets }}")
    doc.add_paragraph("After")

    lines = ["Line A", "Line B", "Line C"]
    _expand_paragraph(para, lines)

    texts = [p.text for p in doc.paragraphs]
    assert "Before" in texts
    assert "After" in texts
    assert "Line A" in texts
    assert "Line B" in texts
    assert "Line C" in texts
    # Original placeholder text must not remain
    assert "{{ bullets }}" not in texts


def test_expand_paragraph_preserves_paragraph_style():
    """New paragraphs must have the same style name as the placeholder paragraph."""
    from jobos.docx_generator import _expand_paragraph

    doc = Document()
    para = doc.add_paragraph("{{ experience_bullets }}", style="List Bullet")
    original_style = para.style.name
    _expand_paragraph(para, ["Bullet one", "Bullet two"])
    # The two new paragraphs should have the same style
    new_paras = [p for p in doc.paragraphs if p.text in ("Bullet one", "Bullet two")]
    assert len(new_paras) == 2
    for p in new_paras:
        assert p.style.name == original_style


def test_expand_paragraph_preserves_run_formatting():
    """Text in expanded paragraphs must inherit bold/italic from the original run."""
    from jobos.docx_generator import _expand_paragraph

    doc = Document()
    para = doc.add_paragraph()
    run = para.add_run("{{ skills }}")
    run.bold = True
    _expand_paragraph(para, ["Skill A", "Skill B"])
    new_paras = [p for p in doc.paragraphs if p.text in ("Skill A", "Skill B")]
    assert len(new_paras) == 2
    for p in new_paras:
        text_runs = [r for r in p.runs if r.text.strip()]
        assert any(r.bold for r in text_runs)


# ── Multiline replacement via _replace_in_doc ─────────────────────────────────

def test_multiline_replacement_creates_separate_paragraphs(tmp_path):
    """Multiline context values must produce separate paragraphs, not soft breaks."""
    from jobos.docx_generator import _replace_in_doc

    doc = Document()
    doc.add_paragraph("Header")
    doc.add_paragraph("{{ skills }}")
    doc.add_paragraph("Footer")

    _replace_in_doc(doc, {"skills": "• Python\n• SQL\n• Excel"})

    texts = [p.text for p in doc.paragraphs]
    assert "• Python" in texts
    assert "• SQL" in texts
    assert "• Excel" in texts
    assert "{{ skills }}" not in texts
    assert "Header" in texts
    assert "Footer" in texts


def test_single_line_replacement_stays_in_one_paragraph(tmp_path):
    """Single-line context values must not expand into multiple paragraphs."""
    from jobos.docx_generator import _replace_in_doc

    doc = Document()
    doc.add_paragraph("{{ name }}")
    _replace_in_doc(doc, {"name": "Alice Smith"})
    texts = [p.text for p in doc.paragraphs]
    assert texts.count("Alice Smith") == 1


# ── Table cell placeholder replacement ────────────────────────────────────────

def test_placeholder_in_table_cell_is_replaced(tmp_path):
    """Placeholders inside table cells must be filled."""
    from jobos.docx_generator import _replace_in_doc

    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).paragraphs[0].add_run("{{ name }}")
    table.cell(0, 1).paragraphs[0].add_run("Static text")
    table.cell(1, 0).paragraphs[0].add_run("{{ role }}")

    _replace_in_doc(doc, {"name": "Alice", "role": "Analyst"})

    assert table.cell(0, 0).text == "Alice"
    assert table.cell(0, 1).text == "Static text"
    assert table.cell(1, 0).text == "Analyst"


def test_multiline_in_table_cell_expands(tmp_path):
    """Multiline values in table cells must expand into multiple paragraphs."""
    from jobos.docx_generator import _replace_in_doc

    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).paragraphs[0].add_run("{{ skills }}")

    _replace_in_doc(doc, {"skills": "• Skill A\n• Skill B"})

    cell_texts = [p.text for p in table.cell(0, 0).paragraphs]
    assert "• Skill A" in cell_texts
    assert "• Skill B" in cell_texts


# ── diagnose_template ──────────────────────────────────────────────────────────

def test_diagnose_template_finds_body_placeholders(tmp_path):
    from jobos.docx_generator import diagnose_template, CV_PLACEHOLDERS

    doc = Document()
    for ph in CV_PLACEHOLDERS:
        doc.add_paragraph(f"{{{{ {ph} }}}}")
    path = tmp_path / "cv.docx"
    doc.save(str(path))

    result = diagnose_template(path, CV_PLACEHOLDERS)
    assert result["is_valid"] is True
    assert result["missing"] == []
    assert set(result["found_body"]) == set(CV_PLACEHOLDERS)
    assert result["found_table"] == []


def test_diagnose_template_finds_table_placeholders(tmp_path):
    from jobos.docx_generator import diagnose_template

    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).paragraphs[0].add_run("{{ name }}")
    table.cell(0, 1).paragraphs[0].add_run("{{ contact_line }}")
    path = tmp_path / "table_cv.docx"
    doc.save(str(path))

    result = diagnose_template(path, ["name", "contact_line", "profile_summary"])
    assert "name" in result["found_table"]
    assert "contact_line" in result["found_table"]
    assert "profile_summary" in result["missing"]
    assert result["is_valid"] is False


def test_diagnose_template_reports_all_missing(tmp_path):
    from jobos.docx_generator import diagnose_template, CV_PLACEHOLDERS

    doc = Document()
    doc.add_paragraph("No placeholders here at all.")
    path = tmp_path / "empty.docx"
    doc.save(str(path))

    result = diagnose_template(path, CV_PLACEHOLDERS)
    assert result["is_valid"] is False
    assert len(result["missing"]) == len(CV_PLACEHOLDERS)
    assert result["all_found"] == []


def test_diagnose_template_mixed_body_and_table(tmp_path):
    from jobos.docx_generator import diagnose_template

    doc = Document()
    doc.add_paragraph("{{ name }}")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).paragraphs[0].add_run("{{ role }}")
    path = tmp_path / "mixed.docx"
    doc.save(str(path))

    result = diagnose_template(path, ["name", "role", "missing_one"])
    assert "name" in result["found_body"]
    assert "role" in result["found_table"]
    assert "missing_one" in result["missing"]
    assert result["is_valid"] is False
